from collections import OrderedDict
from datetime import datetime
import io
import os
import re

from docx import Document
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Cm
import pandas as pd
from PIL import Image, ImageOps
import streamlit as st


# Resim işleme ve otomatik yön/boyut ayarlama fonksiyonu
def process_and_get_image(doc, uploaded_file, width_cm=6.5, height_cm=5.0):
    if uploaded_file is None:
        return ""
    try:
        img = Image.open(uploaded_file)
        img = ImageOps.exif_transpose(img)
        img.thumbnail((1200, 1200))
        img_byte_arr = io.BytesIO()
        img_format = img.format if img.format else "JPEG"
        img.save(img_byte_arr, format=img_format, quality=85)
        img_byte_arr.seek(0)
        return InlineImage(
            doc,
            img_byte_arr,
            width=Cm(width_cm),
            height=Cm(height_cm),
        )
    except Exception:
        return ""


# Alınan yerleri gruplayıp sayılarını hesaplayan yardımcı fonksiyon
def generate_bolum_summary(samples):
    place_counts = OrderedDict()
    for s in samples:
        yer = s["yer"] if s["yer"] and s["yer"] != "-" else "Belirtilmedi"
        place_counts[yer] = place_counts.get(yer, 0) + 1

    bolum_summary = []
    for yer, sayi in place_counts.items():
        bolum_summary.append({"yer": yer, "sayi": sayi})
    return bolum_summary


# Tutanağın üst bilgi ve numune tablosunu okuyan fonksiyon
def parse_asbest_tutanak(file):
    df_raw = pd.read_excel(file, header=None)

    info = {
        "musteri_adi": "ABC İnşaat",
        "adres": "-",
        "pafta": "-",
        "ada": "-",
        "parsel": "-",
        "numune_tarihi": datetime.now().strftime("%d.%m.%Y"),
        "teklif_no": "26-08-5191",
        "telefon": "-",
    }

    for idx in range(min(10, len(df_raw))):
        row_values = [str(x) for x in df_raw.iloc[idx].values if pd.notna(x)]
        row_text = " ".join(row_values)

        if "Talep Numarası" in row_text:
            if idx + 1 < len(df_raw):
                val = str(df_raw.iloc[idx + 1].values[0])
                if val and val != "nan":
                    info["teklif_no"] = val.strip()

        if "Firma Adı:" in row_text:
            m = re.search(r"Firma Adı:\s*(.*?)(?:Telefon|$)", row_text)
            if m and m.group(1).strip():
                info["musteri_adi"] = m.group(1).strip()

        if "Telefon Numarası:" in row_text:
            m = re.search(r"Telefon Numarası:\s*(.*)", row_text)
            if m and m.group(1).strip():
                info["telefon"] = m.group(1).strip()

        if "Firma Adresi:" in row_text:
            m = re.search(r"Firma Adresi:\s*(.*)", row_text)
            if m and m.group(1).strip():
                info["adres"] = m.group(1).strip()

        if "Pafta No:" in row_text or "Parsel No:" in row_text:
            p = re.search(r"Pafta\s*No:\s*([^\s|]*)(?=\s*Ada|$)", row_text, re.IGNORECASE)
            a = re.search(r"Ada\s*No:\s*([^\s|]*)(?=\s*Parsel|$)", row_text, re.IGNORECASE)
            pr = re.search(r"Parsel\s*No:\s*([^\s|]*)(?=$)", row_text, re.IGNORECASE)

            if p and p.group(1).strip():
                info["pafta"] = p.group(1).strip()
            if a and a.group(1).strip():
                info["ada"] = a.group(1).strip()
            if pr and pr.group(1).strip():
                info["parsel"] = pr.group(1).strip()

        if "Tarih" in row_text:
            for cell in row_values:
                if re.match(r"\d{2}\.\d{2}\.\d{4}", cell):
                    info["numune_tarihi"] = cell

    samples = []
    for idx in range(len(df_raw)):
        row = df_raw.iloc[idx]
        row_str = " ".join([str(x) for x in row.values if pd.notna(x)])

        code_match = re.search(r"NK\.\d+\.\d+-\d+", row_str)
        if code_match:
            code = code_match.group(0)
            non_empty = [
                str(x).strip()
                for x in row.values
                if pd.notna(x) and str(x).strip() != ""
            ]

            if len(non_empty) >= 3 and any(
                k in non_empty[1] for k in ["NK.", "NK"]
            ):
                kat = non_empty[1] if len(non_empty) > 1 else "1. Kat"
                tur = non_empty[2] if len(non_empty) > 2 else "Beton / Sıva"
                yer = non_empty[3] if len(non_empty) > 3 else "-"
                yontem = non_empty[4] if len(non_empty) > 4 else "-"
                strateji = non_empty[5] if len(non_empty) > 5 else "-"

                samples.append(
                    {
                        "kod": code,
                        "kat": kat,
                        "tur": tur,
                        "yer": yer,
                        "yontem": yontem,
                        "strateji": strateji,
                    }
                )

    return info, samples


# ANA MODÜL FONKSİYONU
def render_asbest_module():
    st.title("📋 Asbest Katı Numune Analiz Raporu Oluşturucu")

    # 1. Rapor Numarası ve Doğrulama Numarası En Üstte ve Sabit
    st.markdown("### 🔢 Rapor ve Belge Bilgileri")
    col_n1, col_n2 = st.columns(2)
    with col_n1:
        user_rapor_no = st.text_input("Rapor Numarası", value="ARK.26.4861", key="input_rapor_no")
    with col_n2:
        user_dogrulama_no = st.text_input("Doğrulama / Belge Numarası", value="328", key="input_dogrulama_no")

    st.markdown("---")

    uploaded_file = st.file_uploader(
        "Numune Tutanağı Excel Dosyasını Yükleyin", type=["xlsx", "xls"]
    )

    if uploaded_file is not None:
        info, samples = parse_asbest_tutanak(uploaded_file)
        st.success(
            f"Tutanak başarıyla okundu! Toplam **{len(samples)}** adet numune tespit edildi."
        )

        st.markdown("---")
        st.subheader("📅 Genel Bilgiler ve Tarih Ayarları")

        bugun_tarih = datetime.now().strftime("%d.%m.%Y")

        col_m1, col_m2 = st.columns(2)
        with col_m1:
            musteri_adi = st.text_input(
                "Müşteri / Mal Sahibi:", value=info["musteri_adi"]
            )
            adres = st.text_input("Adres:", value=info["adres"])
            teklif_no = st.text_input("Teklif Numarası:", value=info["teklif_no"])
            numune_tarihi = st.text_input(
                "Numune Alma Tarihi (Tutanaktan):", value=info["numune_tarihi"]
            )
        with col_m2:
            pafta_ada_parsel = (
                f"{info['pafta']} / {info['ada']} / {info['parsel']}"
            )
            st.info(f"**Pafta / Ada / Parsel:** {pafta_ada_parsel}")
            rapor_tarihi = st.text_input(
                "Rapor Oluşturulma / Yayın Tarihi:", value=bugun_tarih
            )

        st.markdown("---")
        st.subheader("👨‍💼 Personel Seçimi")

        numune_nezaret_listesi = [
            "Abdul Samed DEĞİRMENCİ",
            "Emir UÇARLI",
            "Ali Kemal DEĞİRMENCİ",
            "Burak BAYRAKTAR",
            "Doğucan TAŞTAN",
            "Emre Can İNEGAZİLİ",
            "Gözde CANİK",
            "Furkan TEMİZ",
            "İsmail AYDIN",
            "Ogün KAN",
            "Muharrem YAŞAR",
        ]

        deney_sorumlusu_listesi = [
            "Gizem DEMİR",
            "Edanur KESGİN",
            "Ali Kemal DEĞİRMENCİ",
        ]

        col1, col2, col3 = st.columns(3)
        with col1:
            numune_alan = st.selectbox(
                "Numune Alan Personel:", options=numune_nezaret_listesi
            )
        with col2:
            nezaret_eden = st.selectbox(
                "Nezaret Eden Personel:", options=numune_nezaret_listesi
            )
        with col3:
            deney_sorumlusu = st.selectbox(
                "Deney Sorumlusu (İmza Yetkilisi):",
                options=deney_sorumlusu_listesi,
            )

        st.markdown("---")
        st.subheader("🖼️ Fotoğraf Yükleme Seçeneği")
        foto_secenegi = st.radio(
            "Rapor fotoğraflarını şimdi yüklemek ister misiniz?",
            [
                "Fotoğrafları Yükleme (Sonradan Word üzerinde eklenecek)",
                "Fotoğrafları Şimdi Yükle",
            ],
            horizontal=True,
        )

        bina_foto_on = None
        bina_foto_yan = None
        bina_foto_arka = None
        numune_fotolari = {}

        if foto_secenegi == "Fotoğrafları Şimdi Yükle":
            st.markdown("##### 🏢 Bina Dış Görünüş Fotoğrafları")
            b_col1, b_col2 = st.columns(2)
            with b_col1:
                bina_foto_on = st.file_uploader("Ön Cephe Fotoğrafı", type=["jpg", "jpeg", "png"], key="bina_on")
            with b_col2:
                bina_foto_yan = st.file_uploader("Yan Cephe Fotoğrafı", type=["jpg", "jpeg", "png"], key="bina_yan")
            bina_foto_arka = st.file_uploader("Arka Cephe Fotoğrafı", type=["jpg", "jpeg", "png"], key="bina_arka")

        st.markdown("---")
        st.subheader("🧪 Numune Sonuçları ve Bilgileri")

        numuneler = []

        for index, s in enumerate(samples):
            n_kodu = s["kod"]
            m_turu = s["tur"]

            st.markdown(
                f"**Numune {index+1} | Kod:** `{n_kodu}` | **Kat:** `{s['kat']}` | **Yer:** `{s['yer']}` | **Malzeme:** `{m_turu}`"
            )

            c1, c2 = st.columns([1, 2])
            with c1:
                asbest_durumu = st.radio(
                    f"Asbest Durumu ({n_kodu})",
                    ["Yok", "Var"],
                    horizontal=True,
                    key=f"asbest_durum_{index}",
                )
            with c2:
                if asbest_durumu == "Var":
                    asbest_turu = st.text_input(
                        "Tespit Edilen Asbest Türü:", key=f"asbest_tur_{index}"
                    )
                    sonuc_metni = (
                        f"Asbest tespit edilmiştir ({asbest_turu})"
                        if asbest_turu
                        else "Asbest tespit edilmiştir"
                    )
                else:
                    sonuc_metni = "Asbest tespit edilmedi"

            # Her numune için Uzak, Yakın ve Poşetli 3'lü fotoğraf sütunu
            if foto_secenegi == "Fotoğrafları Şimdi Yükle":
                f_c1, f_c2, f_c3 = st.columns(3)
                with f_c1:
                    numune_fotolari[f"uzak_{index}"] = st.file_uploader(f"Uzak Fotoğraf ({n_kodu})", type=["jpg", "jpeg", "png"], key=f"uzak_{index}")
                with f_c2:
                    numune_fotolari[f"yakin_{index}"] = st.file_uploader(f"Yakın Fotoğraf ({n_kodu})", type=["jpg", "jpeg", "png"], key=f"yakin_{index}")
                with f_c3:
                    numune_fotolari[f"posetli_{index}"] = st.file_uploader(f"Poşetli Fotoğraf ({n_kodu})", type=["jpg", "jpeg", "png"], key=f"pos_{index}")
                st.markdown("---")

            on_islem = (
                "Asitle Muamele" if "marley" in m_turu.lower() else "Parçalama"
            )

            numuneler.append(
                {
                    "sira": index + 1,
                    "tarih": numune_tarihi,
                    "kod": n_kodu,
                    "kat": s["kat"],
                    "tur": m_turu,
                    "yer": s["yer"],
                    "yontem": s["yontem"],
                    "strateji": s["strateji"],
                    "homojenite": "Homojen",
                    "onislem": on_islem,
                    "sonuc": sonuc_metni,
                }
            )

        st.markdown("---")
        if st.button("🚀 Word Raporunu Oluştur ve İndir", type="primary"):
            try:
                base_dir = os.path.dirname(
                    os.path.dirname(os.path.abspath(__file__))
                )
                template_path = os.path.join(
                    base_dir, "templates", "sablon.docx"
                )
                temp_path = os.path.join(base_dir, "gecici_rapor.docx")
                output_path = os.path.join(base_dir, f"cikis_asbest_raporu_{user_rapor_no}.docx")

                tpl = DocxTemplate(template_path)

                context = {
                    "musteri_adi": musteri_adi,
                    "adres": adres,
                    "teklif_no": teklif_no,
                    "pafta": info["pafta"],
                    "ada": info["ada"],
                    "parsel": info["parsel"],
                    "numune_tarihi": numune_tarihi,
                    "rapor_tarihi": rapor_tarihi,
                    "numune_alan": numune_alan,
                    "nezaret_eden": nezaret_eden,
                    "deney_sorumlusu": deney_sorumlusu,
                    "rapor_no": user_rapor_no,
                    "dogrulama_no": user_dogrulama_no,
                    "samples": samples,
                    "bolum_listesi": generate_bolum_summary(samples),
                }

                if foto_secenegi == "Fotoğrafları Şimdi Yükle":
                    context["bina_foto_on"] = process_and_get_image(tpl, bina_foto_on, width_cm=7.0, height_cm=8.0)
                    context["bina_foto_yan"] = process_and_get_image(tpl, bina_foto_yan, width_cm=7.0, height_cm=8.0)
                    context["bina_foto_arka"] = process_and_get_image(tpl, bina_foto_arka, width_cm=14.0, height_cm=8.0)
                else:
                    context["bina_foto_on"] = ""
                    context["bina_foto_yan"] = ""
                    context["bina_foto_arka"] = ""

                tpl.render(context)
                tpl.save(temp_path)

                doc = Document(temp_path)

                # 1. Analiz Sonuçları Tablosunu Doldur
                target_table = None
                for tbl in doc.tables:
                    if len(tbl.columns) == 10:
                        target_table = tbl
                        break
                if target_table is None:
                    target_table = doc.tables[2]

                while len(target_table.rows) > 2:
                    r = target_table.rows[1]._tr
                    r.getparent().remove(r)

                footer_row = target_table.rows[-1]

                for n in numuneler:
                    new_tr = target_table.add_row()._tr
                    footer_row._tr.addprevious(new_tr)

                    new_row_cells = target_table.rows[-2].cells
                    veriler = [
                        str(n["sira"]),
                        str(n["tarih"]),
                        str(n["kod"]),
                        str(n["tur"]),
                        str(n["yer"]),
                        str(n["yontem"]),
                        str(n["strateji"]),
                        str(n["homojenite"]),
                        str(n["onislem"]),
                        str(n["sonuc"]),
                    ]
                    for i, val in enumerate(veriler):
                        if i < len(new_row_cells):
                            new_row_cells[i].text = val

                    # 2. Numune Fotoğrafları Tablosunu Yatay Sütunlar Halinde (Şablona Uygun) Doldur
                    foto_table = None
                    for tbl in doc.tables:
                        if len(tbl.columns) > 4:  # Çok sütunlu matris fotoğraf tablosu
                            foto_table = tbl
                            break

                    if foto_table and foto_secenegi == "Fotoğrafları Şimdi Yükle":
                        for index, s in enumerate(samples):
                            # Şablonda her numune sırasıyla sütunlara denk geliyor (Genellikle 2. sütundan başlar)
                            target_col_idx = index + 2  
                        
                            if target_col_idx < len(foto_table.columns):
                                img_keys = [f"uzak_{index}", f"yakin_{index}", f"posetli_{index}"]
                                # Satır offsetleri şablondaki yerleşim sırasına göredir (Uzak, Yakın, Poşetli)
                                for row_offset, key in enumerate(img_keys, start=1):
                                    uploaded_img = numune_fotolari.get(key)
                                    if uploaded_img is not None and row_offset < len(foto_table.rows):
                                        cell = foto_table.cell(row_offset, target_col_idx)
                                        p = cell.paragraphs[0]
                                        p.text = ""
                                        run = p.add_run()
                                        inline_img = process_and_get_image(tpl, uploaded_img, width_cm=3.0, height_cm=3.5)
                                        if inline_img:
                                            run._r.append(inline_img._inline)

                with open(output_path, "rb") as file:
                    st.download_button(
                        label="📥 Oluşturulan Raporu İndir (.docx)",
                        data=file,
                        file_name=f"Asbest_Analiz_Raporu_{user_rapor_no}.docx",
                        mime=(
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        ),
                    )
            except Exception as e:
                st.error(f"Rapor oluşturulurken hata meydana geldi: {e}")
