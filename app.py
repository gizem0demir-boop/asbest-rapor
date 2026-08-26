from datetime import datetime
import os
import io
import re
import time
from collections import OrderedDict
from docx import Document
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from PIL import Image, ImageOps
import pandas as pd
import streamlit as st

# Selenium importları (Task Pano otomasyonu için)
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
import zipfile

# Sayfa Yapılandırması
st.set_page_config(
    page_title="Asya Asbest & Atık Yönetim Sistemi",
    page_icon="🧪",
    layout="wide",
)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ==========================================
# YARDIMCI FONKSİYONLAR
# ==========================================

def read_tutanak_details(tutanak_path):
    try:
        df = pd.read_excel(tutanak_path, sheet_name="Table 1", header=None)
    except:
        xls = pd.ExcelFile(tutanak_path)
        df = pd.read_excel(xls, sheet_name=xls.sheet_names[0], header=None)

    raw_firma = str(df.iloc[4, 0]) if pd.notna(df.iloc[4, 0]) else ""
    musteri_adi = raw_firma.replace("Firma Adı:", "").strip()

    raw_adres = str(df.iloc[5, 0]) if pd.notna(df.iloc[5, 0]) else ""
    adres = raw_adres.replace("Firma Adresi:", "").strip()

    raw_pafta = str(df.iloc[6, 0]) if pd.notna(df.iloc[6, 0]) else "-"
    raw_ada = str(df.iloc[6, 4]) if pd.notna(df.iloc[6, 4]) else "-"
    raw_parsel = str(df.iloc[6, 8]) if pd.notna(df.iloc[6, 8]) else "-"

    pafta = raw_pafta.replace("Pafta No:", "").strip() or "-"
    ada = raw_ada.replace("Ada No:", "").strip() or "-"
    parsel = raw_parsel.replace("Parsel No:", "").strip() or "-"
    pafta_ada_parsel = f"{pafta} / {ada} / {parsel}"

    return {
        "musteri_adi": musteri_adi,
        "MUSTERI_ADI": musteri_adi,
        "adres": adres,
        "ADRES": adres,
        "pafta": pafta,
        "ada": ada,
        "parsel": parsel,
        "pafta_ada_parsel": pafta_ada_parsel,
    }

def process_and_get_image(doc, uploaded_file, width_cm=6.5, height_cm=5.0):
    if uploaded_file is None:
        return ""
    try:
        img = Image.open(uploaded_file)
        img = ImageOps.exif_transpose(img)
        img.thumbnail((1200, 1200))
        img_byte_arr = io.BytesIO()
        img_format = img.format if img.format else 'JPEG'
        img.save(img_byte_arr, format=img_format, quality=85)
        img_byte_arr.seek(0)
        return InlineImage(doc, img_byte_arr, width=Mm(width_cm * 10), height=Mm(height_cm * 10))
    except Exception:
        return ""

def generate_bolum_summary(samples):
    place_counts = OrderedDict()
    for s in samples:
        yer = s['yer'] if s['yer'] and s['yer'] != '-' else 'Belirtilmedi'
        place_counts[yer] = place_counts.get(yer, 0) + 1
    return [{'yer': yer, 'sayi': sayi} for yer, sayi in place_counts.items()]

def parse_asbest_tutanak(file):
    df_raw = pd.read_excel(file, header=None)
    info = {
        'musteri_adi': 'ABC İnşaat', 'adres': '-', 'pafta': '-', 
        'ada': '-', 'parsel': '-', 'numune_tarihi': '20.08.2026', 
        'teklif_no': '26-08-5191', 'telefon': '-'
    }
    for idx in range(min(10, len(df_raw))):
        row_values = [str(x) for x in df_raw.iloc[idx].values if pd.notna(x)]
        row_text = " ".join(row_values)
        if "Talep Numarası" in row_text and idx + 1 < len(df_raw):
            val = str(df_raw.iloc[idx+1].values[0])
            if val and val != "nan": info['teklif_no'] = val.strip()
        if "Firma Adı:" in row_text:
            m = re.search(r'Firma Adı:\s*(.*?)(?:Telefon|$)', row_text)
            if m and m.group(1).strip(): info['musteri_adi'] = m.group(1).strip()
        if "Firma Adresi:" in row_text:
            m = re.search(r'Firma Adresi:\s*(.*)', row_text)
            if m and m.group(1).strip(): info['adres'] = m.group(1).strip()

    samples = []
    for idx in range(len(df_raw)):
        row = df_raw.iloc[idx]
        row_str = " ".join([str(x) for x in row.values if pd.notna(x)])
        code_match = re.search(r'NK\.\d+\.\d+-\d+', row_str)
        if code_match:
            code = code_match.group(0)
            non_empty = [str(x).strip() for x in row.values if pd.notna(x) and str(x).strip() != '']
            if len(non_empty) >= 3 and any(k in non_empty[1] for k in ['NK.', 'NK']):
                samples.append({
                    'kod': code, 'tur': non_empty[2] if len(non_empty) > 2 else "Beton / Sıva",
                    'yer': non_empty[3] if len(non_empty) > 3 else "-",
                    'yontem': non_empty[4] if len(non_empty) > 4 else "-",
                    'strateji': non_empty[5] if len(non_empty) > 5 else "-"
                })
    return info, samples

def create_asbest_report(info, numuneler, foto_secenegi, bina_foto, numune_fotolari, samples):
    tpl = DocxTemplate("sablon.docx")
    
    context = {
        "musteri_adi": info['musteri_adi'],
        "adres": info['adres'],
        "teklif_no": info['teklif_no'],
        "pafta": info['pafta'],
        "ada": info['ada'],
        "parsel": info['parsel'],
        "numune_tarihi": info['numune_tarihi'],
        "rapor_tarihi": info['rapor_tarihi'],
        "numune_alan": info['numune_alan'],
        "nezaret_eden": info['nezaret_eden'],
        "deney_sorumlusu": info['deney_sorumlusu'],
        "bolum_listesi": generate_bolum_summary(samples)
    }

    if foto_secenegi == "Fotoğrafları Şimdi Yükle":
        context["bina_foto"] = process_and_get_image(tpl, bina_foto, width_cm=8.0, height_cm=6.0)
        for index, s in enumerate(samples):
            n_kodu = s['kod']
            uploaded_img = numune_fotolari.get(n_kodu)
            img_obj = process_and_get_image(tpl, uploaded_img, width_cm=6.5, height_cm=5.0)
            context[f"foto_{index+1}"] = img_obj
    else:
        context["bina_foto"] = ""
        for index in range(len(samples)):
            context[f"foto_{index+1}"] = ""

    tpl.render(context)
    temp_path = "gecici_rapor.docx"
    tpl.save(temp_path)
    
    doc = Document(temp_path)
    target_table = None
    for tbl in doc.tables:
        if len(tbl.columns) == 10:
            target_table = tbl
            break
    if target_table is None and len(doc.tables) > 2:
        target_table = doc.tables[2]
    
    if target_table and len(target_table.rows) > 2:
        while len(target_table.rows) > 2:
            r = target_table.rows[1]._tr
            r.getparent().remove(r)

    if target_table:
        footer_row = target_table.rows[-1]
        for n in numuneler:
            new_tr = target_table.add_row()._tr
            footer_row._tr.addprevious(new_tr)
            
            new_row_cells = target_table.rows[-2].cells
            veriler = [
                str(n['sira']), str(n['tarih']), str(n['kod']), str(n['tur']),
                str(n['yer']), str(n['yontem']), str(n['strateji']),
                str(n['homojenite']), str(n['onislem']), str(n['sonuc'])
            ]
            for i, val in enumerate(veriler):
                if i < len(new_row_cells):
                    new_row_cells[i].text = val

    output_path = "cikis_asbest_raporu.docx"
    doc.save(output_path)
    return output_path

# Selenium Task Pano Fonksiyonu
def task_pano_dosyalari_indir(email, password, task_url):
    download_dir = os.path.abspath("task_downloads")
    os.makedirs(download_dir, exist_ok=True)

    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    
    prefs = {
        "download.default_directory": download_dir,
        "download.prompt_for_download": False,
        "directory_upgrade": True
    }
    chrome_options.add_experimental_option("prefs", prefs)

    driver = webdriver.Chrome(options=chrome_options)
    try:
        driver.get("https://app.taskpano.com/login")
        time.sleep(2)

        email_input = driver.find_element(By.NAME, "email")
        password_input = driver.find_element(By.NAME, "password")
        
        email_input.send_keys(email)
        password_input.send_keys(password)
        
        login_btn = driver.find_element(By.XPATH, "//button[@type='submit']")
        login_btn.click()
        time.sleep(3)

        driver.get(task_url)
        time.sleep(4)

        download_links = driver.find_elements(By.PARTIAL_LINK_TEXT, "İndir")
        for link in download_links:
            try:
                link.click()
                time.sleep(1)
            except:
                pass

        time.sleep(5)
        
        zip_path = os.path.abspath("task_dosyalar.zip")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(download_dir):
                for file in files:
                    zipf.write(os.path.join(root, file), file)

        return zip_path
    except Exception as e:
        raise Exception(f"Task Pano otomasyon hatası: {str(e)}")
    finally:
        driver.quit()

# ==========================================
# YAN MENÜ (SIDEBAR)
# ==========================================
with st.sidebar:
    st.image("https://img.icons8.com/color/96/experimental-copy.png", width=80)
    st.markdown("### 🔬 Laboratuvar Modülü")
    st.write("ASYA Asbest Danışmanlık ve Laboratuvar Hizmetleri Otomasyon Paneli")
    st.markdown("---")
    
    rapor_turu = st.selectbox(
        "📋 İşlem / Rapor Türü Seçin:",
        [
            "-- Seçiniz --",
            "🧪 Asbest Tür Tayini Raporu",
            "💨 Toz Raporu",
            "♻️ AYP (Atık Yönetim Planı) Raporu",
        ],
    )
    
    st.markdown("---")
    st.markdown("### ⚙️ Task Pano Ayarları")
    tp_email = st.text_input("Task Pano E-posta:", type="default")
    tp_sifre = st.text_input("Task Pano Şifre:", type="password")

# ==========================================
# ANA EKRAN & MODÜL YÖNLENDİRMELERİ
# ==========================================
if rapor_turu == "-- Seçiniz --":
    st.title("🧪 Asbest ve Atık Yönetim Rapor Sistemi")
    st.markdown("---")
    st.warning("⚠️ Lütfen sol menüden oluşturmak istediğiniz **Rapor Türünü** seçin.")

elif rapor_turu == "🧪 Asbest Tür Tayini Raporu":
    st.title("🧪 Asbest Katı Numune Analiz Raporu Oluşturucu")
    st.markdown("---")

    st.markdown("### 🔗 Task Pano Veri ve Dosya İndirici")
    col_tp1, col_tp2 = st.columns([3, 1])
    with col_tp1:
        task_link = st.text_input("Task Pano Görev Linkini Buraya Yapıştırın:", placeholder="https://app.taskpano.com/task/...")
    with col_tp2:
        st.write("")
        st.write("")
        indir_buton = st.button("📥 Dosyaları Çek & İndir")
        
    if indir_buton:
        if not tp_email or not tp_sifre:
            st.error("Lütfen sol menüden Task Pano e-posta ve şifrenizi girin!")
        elif not task_link:
            st.warning("Lütfen geçerli bir Task Pano görev linki yapıştırın.")
        else:
            with st.spinner("Task Pano'ya bağlanılıyor, dosyalar indiriliyor ve arşivleniyor... Lütfen bekleyin."):
                try:
                    zip_dosya_yolu = task_pano_dosyalari_indir(tp_email, tp_sifre, task_link)
                    st.success("Tüm dosyalar başarıyla çekildi ve paketlendi!")
                    with open(zip_dosya_yolu, "rb") as f:
                        st.download_button(
                            label="📦 Arşivlenmiş Dosyaları İndir (.zip)",
                            data=f,
                            file_name="task_pano_dosyalar.zip",
                            mime="application/zip"
                        )
                except Exception as e:
                    st.error(f"İşlem sırasında hata oluştu: {e}")

    st.markdown("---")
    uploaded_file = st.file_uploader("Numune Tutanağı Excel Dosyasını Yükleyin", type=["xlsx", "xls"])

    if uploaded_file is not None:
        info, samples = parse_asbest_tutanak(uploaded_file)
        st.success(f"Tutanak başarıyla okundu! Toplam **{len(samples)}** adet numune tespit edildi.")

        st.markdown("---")
        st.subheader("🏢 Genel Bilgiler ve Tarih Ayarları")
        bugun_tarih = datetime.now().strftime("%d.%m.%Y")
        
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            musteri_adi = st.text_input("Müşteri / Mal Sahibi:", value=info['musteri_adi'])
            adres = st.text_input("Adres:", value=info['adres'])
            teklif_no = st.text_input("Teklif Numarası:", value=info['teklif_no'])
            numune_tarihi = st.text_input("Numune Alma Tarihi (Tutanaktan):", value=info['numune_tarihi'])
        with col_m2:
            pafta_ada_parsel = f"{info['pafta']} / {info['ada']} / {info['parsel']}"
            st.info(f"**Pafta / Ada / Parsel:** {pafta_ada_parsel}")
            rapor_tarihi = st.text_input("Rapor Oluşturulma / Yayın Tarihi:", value=bugun_tarih)

        st.markdown("---")
        st.subheader("👥 Personel Seçimi")
        numune_nezaret_listesi = [
            "Abdul Samed DEĞİRMENCİ", "Emir UÇARLI", "Ali Kemal DEĞİRMENCİ", 
            "Burak BAYRAKTAR", "Doğucan TAŞTAN", "Emre Can İNEGAZİLİ", 
            "Gözde CANİK", "Furkan TEMİZ", "İsmail AYDIN", "Ogün KAN", "Muharrem YAŞAR"
        ]
        deney_sorumlusu_listesi = ["Gizem DEMİR", "Edanur KESGİN", "Ali Kemal DEĞİRMENCİ"]
        
        col1, col2, col3 = st.columns(3)
        with col1:
            numune_alan = st.selectbox("Numune Alan Personel:", options=numune_nezaret_listesi)
        with col2:
            nezaret_eden = st.selectbox("Nezaret Eden Personel:", options=numune_nezaret_listesi)
        with col3:
            deney_sorumlusu = st.selectbox("Deney Sorumlusu (İmza Yetkilisi):", options=deney_sorumlusu_listesi)

        st.markdown("---")
        st.subheader("🖼️ Fotoğraf Yükleme Seçeneği")
        foto_secenegi = st.radio(
            "Rapor fotoğraflarını şimdi yüklemek ister misiniz?",
            ["Fotoğrafları Yükleme (Sonradan Word üzerinde eklenecek)", "Fotoğrafları Şimdi Yükle"],
            horizontal=True
        )
        
        bina_foto = None
        numune_fotolari = {}

        if foto_secenegi == "Fotoğrafları Şimdi Yükle":
            st.markdown("##### 🏢 Bina / Konut Fotoğrafı")
            bina_foto = st.file_uploader("Bina Dış Görünüş Fotoğrafı", type=["jpg", "jpeg", "png"], key="bina_foto_uploader")

        st.markdown("---")
        st.subheader("📋 Numune Sonuçları ve Bilgileri")
        
        numuneler = []
        for index, s in enumerate(samples):
            n_kodu = s['kod']
            m_turu = s['tur']

            st.markdown(f"**Numune {index+1} | Kod:** `{n_kodu}` | **Malzeme:** `{m_turu}`")
            
            c1, c2, c3 = st.columns([1, 1.5, 1.5]) if foto_secenegi == "Fotoğrafları Şimdi Yükle" else st.columns([1, 2])
            with c1:
                asbest_durumu = st.radio(f"Asbest Durumu ({n_kodu})", ["Yok", "Var"], horizontal=True, key=f"asbest_durum_{index}")
            with c2:
                if asbest_durumu == "Var":
                    asbest_turu = st.text_input("Tespit Edilen Asbest Türü:", key=f"asbest_tur_{index}")
                    sonuc_metni = f"Asbest tespit edilmiştir ({asbest_turu})" if asbest_turu else "Asbest tespit edilmiştir"
                else:
                    sonuc_metni = "Asbest tespit edilmedi"
            if foto_secenegi == "Fotoğrafları Şimdi Yükle":
                with c3:
                    numune_fotolari[n_kodu] = st.file_uploader(f"Numune Fotoğrafı ({n_kodu})", type=["jpg", "jpeg", "png"], key=f"foto_upl_{index}")

            on_islem = "Asitle Muamele" if "marley" in m_turu.lower() else "Parçalama"
            numuneler.append({
                "sira": index + 1,
                "tarih": numune_tarihi,
                "kod": n_kodu,
                "tur": m_turu,
                "yer": s['yer'],
                "yontem": s['yontem'],
                "strateji": s['strateji'],
                "homojenite": "Homojen",
                "onislem": on_islem,
                "sonuc": sonuc_metni
            })

        st.markdown("---")
        if st.button("🚀 Word Raporunu Oluştur ve İndir", type="primary"):
            try:
                info_dict = {
                    "musteri_adi": musteri_adi,
                    "adres": adres,
                    "teklif_no": teklif_no,
                    "pafta": info['pafta'],
                    "ada": info['ada'],
                    "parsel": info['parsel'],
                    "numune_tarihi": numune_tarihi,
                    "rapor_tarihi": rapor_tarihi,
                    "numune_alan": numune_alan,
                    "nezaret_eden": nezaret_eden,
                    "deney_sorumlusu": deney_sorumlusu
                }
                
                output_path = create_asbest_report(info_dict, numuneler, foto_secenegi, bina_foto, numune_fotolari, samples)
                st.success("Rapor başarıyla oluşturuldu!")
                
                with open(output_path, "rb") as file:
                    st.download_button(
                        label="📥 Oluşturulan Raporu İndir (.docx)",
                        data=file,
                        file_name=f"Asbest_Analiz_Raporu_{teklif_no}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"Hata: {e}")

elif rapor_turu == "💨 Toz Raporu":
    st.title("💨 Toz Ölçüm Raporu Oluşturucu")
    st.markdown("---")
    tutanak_file = st.file_uploader("📂 Tutanak Dosyası (Excel):", type=["xlsx", "xls"], key="toz_tutanak")

    if tutanak_file:
        try:
            tutanak_path = os.path.join(UPLOAD_FOLDER, tutanak_file.name)
            with open(tutanak_path, "wb") as f:
                f.write(tutanak_file.getbuffer())

            info = read_tutanak_details(tutanak_path)
            st.success("✅ Toz tutanak dosyası başarıyla okundu.")

            if st.button("🚀 Toz Raporunu Oluştur ve İndir", type="primary"):
                if os.path.exists("sablon_toz.docx"):
                    doc = DocxTemplate("sablon_toz.docx")
                    doc.render(info)
                    output_path = os.path.join(UPLOAD_FOLDER, "Toz_Raporu_Cikti.docx")
                    doc.save(output_path)
                    st.success("✅ Toz Raporu başarıyla oluşturuldu!")
                    with open(output_path, "rb") as f:
                        st.download_button("📥 Toz Raporunu İndir (.docx)", f, file_name=f"Toz_Raporu_{info['musteri_adi']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error("❌ 'sablon_toz.docx' bulunamadı!")
        except Exception as e:
            st.error(f"Hata: {e}")

elif rapor_turu == "♻️ AYP (Atık Yönetim Planı) Raporu":
    st.title("♻️ Atık Yönetim Planı (AYP) Rapor Oluşturucu")
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        tutanak_file = st.file_uploader("📂 Tutanak Dosyası:", type=["xlsx", "xls"], key="ayp_tutanak")
    with col2:
        ayp_file = st.file_uploader("📂 AYP Hesaplama Dosyası:", type=["xlsx", "xls"], key="ayp_excel")

    if tutanak_file and ayp_file:
        try:
            tutanak_path = os.path.join(UPLOAD_FOLDER, tutanak_file.name)
            with open(tutanak_path, "wb") as f:
                f.write(tutanak_file.getbuffer())
            info = read_tutanak_details(tutanak_path)

            ayp_path = os.path.join(UPLOAD_FOLDER, ayp_file.name)
            with open(ayp_path, "wb") as f:
                f.write(ayp_file.getbuffer())

            df_sayfa2 = pd.read_excel(ayp_path, sheet_name="Sayfa2")
            genel_toplam = 0
            for _, row in df_sayfa2.iterrows():
                if str(row.iloc[4]).strip().lower() == "toplam":
                    genel_toplam = row.iloc[6]

            bugun_tarihi = datetime.now().strftime("%d.%m.%Y")
            info.update({
                "tarih": bugun_tarihi,
                "rapor_tarihi": bugun_tarihi,
                "genel_toplam_miktar": (genel_toplam if genel_toplam != 0 else 236955.7),
            })

            st.success("✅ Dosyalar başarıyla okundu.")
            if st.button("🚀 AYP Raporunu Oluştur ve İndir", type="primary"):
                if os.path.exists("sablon_ayp.docx"):
                    doc = DocxTemplate("sablon_ayp.docx")
                    doc.render(info)
                    output_path = os.path.join(UPLOAD_FOLDER, "AYP_Raporu_Cikti.docx")
                    doc.save(output_path)
                    st.success("✅ AYP Raporu başarıyla oluşturuldu!")
                    with open(output_path, "rb") as f:
                        st.download_button("📥 AYP Raporunu İndir (.docx)", f, file_name=f"AYP_Raporu_{info['musteri_adi']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error("❌ 'sablon_ayp.docx' bulunamadı!")
        except Exception as e:
            st.error(f"Hata: {e}")
